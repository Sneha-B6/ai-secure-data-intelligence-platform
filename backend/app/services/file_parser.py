"""
File Parser module — extracts text from uploaded files.
Supports: .txt, .log, .json, .pdf, .docx, .doc, .rtf
Never raises — returns empty string on failure.
"""
from io import BytesIO
import json
import re
import logging

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract readable text from uploaded files."""
    filename = filename.lower().strip()

    try:
        # TXT and LOG files
        if filename.endswith(".txt") or filename.endswith(".log"):
            return file_bytes.decode("utf-8", errors="ignore")

        # JSON
        if filename.endswith(".json"):
            return _parse_json(file_bytes)

        # DOCX
        if filename.endswith(".docx"):
            return _parse_docx(file_bytes)

        # PDF
        if filename.endswith(".pdf"):
            return _parse_pdf(file_bytes)

        # DOC / RTF
        if filename.endswith(".doc") or filename.endswith(".rtf"):
            return _parse_doc_rtf(file_bytes)

    except Exception as e:
        logger.error(f"File parsing failed for {filename}: {e}")
        return ""

    logger.warning(f"Unsupported file type: {filename}")
    return ""


def _parse_json(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        data = json.loads(text)
        return json.dumps(data, indent=2)
    except json.JSONDecodeError:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"JSON parsing failed: {e}")
        return ""


def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return ""


def _parse_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            if row:
                                text_parts.append(" | ".join(str(cell or "") for cell in row))
        result = "\n".join(text_parts)
        if not result.strip():
            logger.warning("PDF extraction returned empty — may be image-based")
        return result
    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        return ""


def _parse_doc_rtf(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        text = re.sub(r"\\[a-z]+\d*\s?", " ", text)
        text = re.sub(r"\{[^{}]*\}", "", text)
        text = text.replace("{", "").replace("}", "")
        text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
        text = re.sub(r"\\u\d+\s?\??", "", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s*\n+", "\n", text)
        lines = [line.strip() for line in text.split("\n") if len(line.strip()) > 2]
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"DOC/RTF parsing failed: {e}")
        return ""